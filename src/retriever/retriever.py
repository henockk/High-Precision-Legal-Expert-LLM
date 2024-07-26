import os
from dotenv import find_dotenv, load_dotenv
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma

# Load environment variables
load_dotenv(find_dotenv())
os.environ["LANGCHAIN_API_KEY"] = str(os.getenv("LANGCHAIN_API_KEY"))
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
os.environ["LANGCHAIN_PROJECT"] = "default"

def load_docx_document(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)



def create_retriever(doc_path):
    doc_text = load_docx_document(doc_path)
    docs = [Document(page_content=doc_text)]

    text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=300, chunk_overlap=50
    )
    splits = text_splitter.split_documents(docs)

    vectorstore = Chroma.from_documents(
        documents=splits, embedding=OpenAIEmbeddings()
    )
    return vectorstore.as_retriever()
